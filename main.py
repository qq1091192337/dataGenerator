import os

import docx.document
from docx import Document
from docx.table import *

from docx.enum.table import *
from docx.oxml.shared import OxmlElement, qn
import time
import re
inputRoot='./input'
detectHeader={
    "序号":0,
    "学院":1,
    "班级":2,
    "姓名":3,
    "学号":4,
    "性别":5,
    "年级":6,
    "健康码及转码时间":7,
    "体温":8,
    "检测时间":9,
}

data=[[],[],[],[],[],[],[],[],[],[]]
def preProcessCell(cellText :str,col :int):
    if col ==1 and '土木' in cellText:
        return "土木学院"
    if col ==6:
        cellText= cellText.replace('2020','20')
        cellText= cellText.replace('2019','19')
    # if '土木' in cellText:
    #     return "土木学院"
    return cellText.strip()
def docSortFunc(x):
    match= re.match(r'.+?(\d+)级.+',x)
    if match:
        return int(match.group(1))
    else:
        return 22
if __name__ == '__main__':
    input_list= os.listdir('./input')
    input_list= sorted(list(filter(lambda e:e.split('.')[-1]=='docx',input_list)),key=docSortFunc)
    isInsertHeader=True
    counter = 1
    for e in input_list:
        doc :docx.document.Document = Document(f"{inputRoot}/{e}")  # filename为word文档
        # 读取第1个表格
        tb1: Table = doc.tables[0]
        paragraphs = doc.paragraphs

        # 获取第一个表格的列
        for col in tb1.columns:
            header=""
            for cell in col.cells:
                if cell.text.strip() in detectHeader.keys() :

                    header=cell.text.strip()
                    if isInsertHeader:
                        data[detectHeader[header]].append(preProcessCell(cell.text,detectHeader[header]))
                else:

                    # if preProcessCell(cell.text) == "":
                    #     continue

                    if header=="序号":

                        data[detectHeader[header]].append("")
                        # counter+=1
                    else:
                        data[detectHeader[header]].append(preProcessCell(cell.text,detectHeader[header]))
        isInsertHeader=False



    newDocx =Document()
    now= time.localtime(time.time())

    newDocx.add_heading(f'{now.tm_mon}月{now.tm_mday}日土木学院学生晨（午）检及健康码异常情况汇总表', 1)
    table :Table= newDocx.add_table(rows=len (data[0]),cols=9)

    for i in range(0,len(data)-1):

        for j in range(len (data[0])-1):

            table.columns[i].cells[j].text = data[i][j]
            table.columns[i].cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # table.columns[i].cells[j] = WD_TABLE_ALIGNMENT.CENTER

    newDocx.save(f"./output/{now.tm_mon}月{now.tm_mday}日土木学院学生晨（午）检及健康码异常情况汇总表.docx")

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
